import fitz  # PyMuPDF
from pathlib import Path
from docx import Document
import requests
from io import BytesIO
import easyocr
from PIL import Image

# Initialize EasyOCR reader (lazy loading - only when needed)
_easyocr_reader = None

def get_easyocr_reader():
    """Lazy initialization of EasyOCR reader to avoid loading on import."""
    global _easyocr_reader
    if _easyocr_reader is None:
        _easyocr_reader = easyocr.Reader(['en'], gpu=False)  # Set gpu=True if you have GPU
    return _easyocr_reader


def extract_text(file_path_or_url: str) -> str:
    """
    Detects whether input is a URL or local path → extracts text accordingly.
    """
    
    # CASE 1: URL (Supabase)
    if file_path_or_url.startswith("http"):
        extension = file_path_or_url.split("?")[0].split(".")[-1].lower()

        if extension == "pdf":
            return extract_pdf_from_url(file_path_or_url)

        if extension == "txt":
            return extract_txt_from_url(file_path_or_url)

        if extension == "docx":
            return extract_docx_from_url(file_path_or_url)

        raise ValueError(f"Unsupported URL file type: {extension}")


    
    # CASE 2: LOCAL FILE
    file_path = Path(file_path_or_url)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return extract_pdf_local(file_path)

    if ext == ".txt":
        return extract_txt(file_path)

    if ext == ".docx":
        return extract_docx_local(file_path)

    raise ValueError(f"Unsupported local file type: {ext}")


def is_scanned_pdf(doc: fitz.Document) -> bool:
    """
    Checks if PDF is scanned (image-based) by checking if pages have extractable text.
    Returns True if most pages have little/no text.
    """
    total_pages = len(doc)
    if total_pages == 0:
        return True
    
    pages_with_text = 0
    pages_to_check = min(3, total_pages)  # Check first 3 pages or all if less than 3
    
    for page_num in range(pages_to_check):
        page = doc[page_num]
        text = page.get_text("text").strip()
        if len(text) > 50:  # Threshold: if page has >50 chars, it's text-based
            pages_with_text += 1
    
    # If less than 50% of checked pages have text, consider it scanned
    return pages_with_text < (pages_to_check * 0.5)


def pdf_pages_to_images(doc: fitz.Document) -> list:
    """
    Converts PDF pages to PIL Images using PyMuPDF.
    """
    images = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        # Render page to a pixmap (image) at 300 DPI
        mat = fitz.Matrix(300/72, 300/72)  # 300 DPI scaling
        pix = page.get_pixmap(matrix=mat)
        # Convert to PIL Image
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        images.append(img)
    return images


def extract_pdf_with_ocr_local(path: Path) -> str:
    """
    Extracts text from scanned PDF using EasyOCR.
    Uses PyMuPDF to convert pages to images.
    """
    try:
        reader = get_easyocr_reader()
        doc = fitz.open(str(path))
        
        try:
            # Convert PDF pages to images using PyMuPDF
            images = pdf_pages_to_images(doc)
            text_parts = []
            
            for i, image in enumerate(images):
                # EasyOCR returns list of (bbox, text, confidence)
                results = reader.readtext(image)
                # Extract text from results (index 1 is the text)
                page_text = "\n".join([result[1] for result in results])
                if page_text.strip():  # Only add if there's text
                    text_parts.append(f"--- Page {i+1} ---\n{page_text}")
            
            return "\n\n".join(text_parts) if text_parts else ""
        finally:
            doc.close()
            
    except Exception as e:
        raise ValueError(f"OCR extraction failed: {str(e)}")


def extract_pdf_with_ocr_from_url(url: str) -> str:
    """
    Extracts text from scanned PDF URL using EasyOCR.
    Uses PyMuPDF to convert pages to images (no Poppler needed).
    """
    try:
        reader = get_easyocr_reader()
        response = requests.get(url)
        pdf_bytes = response.content
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        try:
            # Convert PDF pages to images using PyMuPDF
            images = pdf_pages_to_images(doc)
            text_parts = []
            
            for i, image in enumerate(images):
                results = reader.readtext(image)
                page_text = "\n".join([result[1] for result in results])
                if page_text.strip():  # Only add if there's text
                    text_parts.append(f"--- Page {i+1} ---\n{page_text}")
            
            return "\n\n".join(text_parts) if text_parts else ""
        finally:
            doc.close()
            
    except Exception as e:
        raise ValueError(f"OCR extraction from URL failed: {str(e)}")


# PDF EXTRACTION — LOCAL
def extract_pdf_local(path: Path) -> str:
    """
    Extracts text from PDF. Automatically detects if it's scanned or text-based.
    Uses PyMuPDF for text-based PDFs, EasyOCR for scanned PDFs.
    """
    doc = fitz.open(str(path))
    
    try:
        # Try regular text extraction first
        text = "\n".join([page.get_text("text") for page in doc])
        
        # Check if PDF is scanned (image-based)
        # If extracted text is very short or pages have no text, use OCR
        if len(text.strip()) < 100 or is_scanned_pdf(doc):
            # Use OCR for scanned PDFs (doc will be closed inside the function)
            # We need to reopen it or pass the doc object
            doc_path = str(path)
            doc.close()
            return extract_pdf_with_ocr_local(Path(doc_path))
        
        # Return text-based extraction (faster and more accurate for text PDFs)
        return text
        
    finally:
        if not doc.is_closed:
            doc.close()


# PDF EXTRACTION — URL
def extract_pdf_from_url(url: str) -> str:
    """
    Extracts text from PDF URL. Automatically detects if it's scanned or text-based.
    Uses PyMuPDF for text-based PDFs, EasyOCR for scanned PDFs.
    """
    response = requests.get(url)
    pdf_bytes = response.content
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    try:
        # Try regular text extraction first
        text = "\n".join([page.get_text("text") for page in doc])
        
        # Check if PDF is scanned (image-based)
        if len(text.strip()) < 100 or is_scanned_pdf(doc):
            # Use OCR for scanned PDFs (doc will be closed inside the function)
            # We need to pass the URL since we can't reopen from bytes easily
            doc.close()
            return extract_pdf_with_ocr_from_url(url)
        
        # Return text-based extraction
        return text
        
    finally:
        if not doc.is_closed:
            doc.close()


# TXT EXTRACTION — LOCAL
def extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")

# TXT EXTRACTION — URL
def extract_txt_from_url(url: str) -> str:
    response = requests.get(url)
    return response.text

# DOCX EXTRACTION — LOCAL
def extract_docx_local(path: Path) -> str:
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

# DOCX EXTRACTION — URL
def extract_docx_from_url(url: str) -> str:
    response = requests.get(url)
    file_bytes = BytesIO(response.content)
    doc = Document(file_bytes)
    return "\n".join([p.text for p in doc.paragraphs])
